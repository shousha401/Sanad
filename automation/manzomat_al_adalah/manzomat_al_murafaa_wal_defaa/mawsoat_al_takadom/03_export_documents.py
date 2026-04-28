"""Step 03: export full documents for each row from the list export."""

from __future__ import annotations

from pathlib import Path
import csv
import hashlib
import json

from docx import Document


STATE_FILE = Path("automation/.state/mawsoat_al_takadom.json")
LIST_FILE = Path("automation/output/manzomat_al_adalah/manzomat_al_murafaa_wal_defaa/mawsoat_al_takadom/list_export.csv")
DOCS_DIR = Path("automation/output/manzomat_al_adalah/manzomat_al_murafaa_wal_defaa/mawsoat_al_takadom/documents")


def short_hash(value: str) -> str:
    return hashlib.md5(value.encode("utf-8", errors="ignore")).hexdigest()[:10]


def export_documents() -> None:
    if not LIST_FILE.exists():
        raise RuntimeError("Step 02 must run first: missing list export.")

    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    with LIST_FILE.open("r", encoding="utf-8-sig", newline="") as f:
        rows = list(csv.DictReader(f))

    exported = []
    for row in rows:
        title = row["title"]
        filename = f"{row['row_index']}_{short_hash(title)}.docx"
        path = DOCS_DIR / filename

        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph("Exported by 03_export_documents.py")
        doc.save(path)
        exported.append(str(path))

    state = {}
    if STATE_FILE.exists():
        state = json.loads(STATE_FILE.read_text(encoding="utf-8"))
    state["current_step"] = "03_export_documents"
    state["documents_exported"] = exported
    STATE_FILE.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"Exported {len(exported)} document(s) to: {DOCS_DIR}")


if __name__ == "__main__":
    export_documents()
