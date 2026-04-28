from pywinauto import Application
from pathlib import Path
from docx import Document
import pyautogui
import time
import hashlib
import csv
import traceback

OUT_DIR = Path(r"C:\EladalahExport\docs")
LOG_FILE = Path(r"C:\EladalahExport\logs\copytext_pilot_log.csv")

OUT_DIR.mkdir(parents=True, exist_ok=True)
LOG_FILE.parent.mkdir(parents=True, exist_ok=True)

TREE_X_REL = 1540
FIRST_ITEM_Y_REL = 278
TREE_ROW_STEP = 28

MAX_ITEMS = 5

def short_hash(text):
    return hashlib.md5(text.encode("utf-8", errors="ignore")).hexdigest()[:8]

def get_info(control, attr):
    try:
        return getattr(control.element_info, attr, "") or ""
    except Exception:
        return ""

def get_name(control):
    try:
        return control.window_text() or get_info(control, "name")
    except Exception:
        return get_info(control, "name")

def get_eladalah_window():
    app = Application(backend="uia").connect(path="eladalah2025.exe", timeout=10)

    candidates = []
    for w in app.windows():
        try:
            title = w.window_text()
            rect = w.rectangle()
            area = (rect.right - rect.left) * (rect.bottom - rect.top)
            if "العدالة" in title or "2025" in title:
                candidates.append((area, w, rect, title))
        except Exception:
            pass

    if not candidates:
        raise RuntimeError("Could not find Eladalah window.")

    candidates.sort(reverse=True, key=lambda x: x[0])
    _, w, rect, title = candidates[0]

    print(f"Using window: {title}")
    w.set_focus()
    time.sleep(0.5)

    try:
        w.maximize()
        time.sleep(1)
    except Exception:
        pass

    rect = w.rectangle()
    print(f"Rect: {rect}")
    return w, rect

def read_rt_txt(main_window):
    """
    Reads the legal text directly from Eladalah rich text control.
    """
    controls = main_window.descendants()

    for c in controls:
        auto_id = get_info(c, "automation_id")
        class_name = get_info(c, "class_name")

        if auto_id == "rt_txt" or "RichEdit20W" in class_name:
            # Try multiple ways to read the text
            methods = []

            try:
                methods.append(c.window_text())
            except Exception:
                pass

            try:
                methods.append(c.texts()[0])
            except Exception:
                pass

            try:
                methods.append(c.get_value())
            except Exception:
                pass

            try:
                methods.append(c.iface_value.CurrentValue)
            except Exception:
                pass

            for text in methods:
                if text and len(text.strip()) > 20:
                    return text.strip()

    return ""

def save_text_to_docx(text, out_path, title):
    doc = Document()
    doc.add_heading(title, level=1)

    for line in text.splitlines():
        line = line.strip()
        if line:
            doc.add_paragraph(line)

    doc.save(str(out_path))

def main():
    pyautogui.FAILSAFE = True
    pyautogui.PAUSE = 0.25

    main_win, rect = get_eladalah_window()

    tree_x = rect.left + TREE_X_REL
    print(f"Tree click X: {tree_x}")

    with LOG_FILE.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["status", "index", "click_x", "click_y", "chars", "output_file", "error"])

        for i in range(MAX_ITEMS):
            item_y = rect.top + FIRST_ITEM_Y_REL + (i * TREE_ROW_STEP)
            out_path = OUT_DIR / f"directtext-pilot-{i+1:03d}-{short_hash(str(item_y))}.docx"

            print(f"\nProcessing row {i+1}")
            print(f"Click tree: {tree_x}, {item_y}")

            try:
                main_win.set_focus()
                time.sleep(0.3)

                pyautogui.click(tree_x, item_y)
                time.sleep(2)

                text = read_rt_txt(main_win)

                if len(text) < 20:
                    raise RuntimeError(f"Read text too short: {repr(text)}")

                title = text.splitlines()[0][:80] if text.splitlines() else f"Row {i+1}"
                save_text_to_docx(text, out_path, title)

                print(f"Read chars: {len(text)}")
                print(f"Saved: {out_path}")

                writer.writerow(["OK", i + 1, tree_x, item_y, len(text), str(out_path), ""])

            except Exception as e:
                print("FAILED")
                print(traceback.format_exc())
                writer.writerow(["FAILED", i + 1, tree_x, item_y, 0, str(out_path), str(e)])

            f.flush()
            time.sleep(1)

    print("\nFinished.")
    print(f"Docs: {OUT_DIR}")
    print(f"Log: {LOG_FILE}")

if __name__ == "__main__":
    main()