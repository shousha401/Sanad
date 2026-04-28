from pywinauto import Application
from pywinauto.keyboard import send_keys
from pathlib import Path
import pyautogui
import time
import hashlib
import csv
import traceback
import re
import ctypes

BASE_DIR = Path(r"C:\EladalahExport")
DOCX_DIR = BASE_DIR / "docs_full"
TXT_DIR = BASE_DIR / "txt_full"
LOG_FILE = BASE_DIR / "logs" / "full_export_log.csv"

DOCX_DIR.mkdir(parents=True, exist_ok=True)
TXT_DIR.mkdir(parents=True, exist_ok=True)
LOG_FILE.parent.mkdir(parents=True, exist_ok=True)

MIN_CHARS = 40
MAX_PAGES = 50
MAX_ITEMS_PER_PAGE = None
STOP_AFTER_STAGNANT_PAGES = 2
NAV_WAIT_SECONDS = 1.1
COM_ERROR = getattr(ctypes, "COMError", Exception)
AUTO_EXPAND = False
MANUAL_EXPANDED_MODE = True
STOP_ON_REPEATED_PAGE_SIGNATURE = True
STOP_AFTER_REPEAT_COUNT = 1
SAVE_TXT = True
SAVE_DOCX = False
FAST_MODE = True
CLICK_WAIT_SECONDS = 0.45
READ_WAIT_SECONDS = 0.25
MOVE_WAIT_SECONDS = 0.6

def text_hash(text):
    return hashlib.md5(text.encode("utf-8", errors="ignore")).hexdigest()

def short_hash(text):
    return text_hash(text)[:10]

def clean_filename(text, max_len=80):
    text = re.sub(r'[\\/:*?"<>|]', " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return (text or "untitled")[:max_len]

def get_info(control, attr):
    try:
        return getattr(control.element_info, attr, "") or ""
    except Exception:
        return ""

def get_name(control):
    try:
        return (control.window_text() or get_info(control, "name") or "").strip()
    except Exception:
        return (get_info(control, "name") or "").strip()

def get_eladalah_window():
    app = Application(backend="uia").connect(path="eladalah2025.exe", timeout=10)

    candidates = []
    for w in app.windows():
        try:
            title = w.window_text()
            rect = w.rectangle()
            area = (rect.right - rect.left) * (rect.bottom - rect.top)
            if "العدالة" in title or "2025" in title:
                candidates.append((area, w, rect, title))
        except Exception:
            pass

    if not candidates:
        raise RuntimeError("Could not find Eladalah window.")

    candidates.sort(reverse=True, key=lambda x: x[0])
    _, w, rect, title = candidates[0]

    print(f"Using window: {title}")
    safe_focus_window(w, rect)
    time.sleep(0.5)

    try:
        w.maximize()
        time.sleep(1)
    except Exception:
        pass

    rect = w.rectangle()
    print(f"Rect: {rect}")
    return w, rect

def safe_focus_window(main_win, rect=None):
    try:
        main_win.set_focus()
        return True, "set_focus"
    except (COM_ERROR, PermissionError, OSError):
        pass
    except Exception as e:
        if "Access is denied" not in str(e):
            pass

    try:
        if rect is None:
            rect = main_win.rectangle()
        x = max(rect.left + 20, (rect.left + rect.right) // 2)
        y = max(rect.top + 20, rect.top + 40)
        pyautogui.click(x, y)
        return True, "window_click_fallback"
    except Exception:
        return False, "focus_failed"

def find_rt_txt(main_window):
    for c in main_window.descendants():
        auto_id = get_info(c, "automation_id")
        class_name = get_info(c, "class_name")
        if auto_id == "rt_txt" or "RichEdit20W" in class_name:
            return c
    return None

def read_rt_txt(rt_ctrl):
    if rt_ctrl is None:
        return ""

    tries = []
    try:
        tries.append(rt_ctrl.window_text())
    except Exception:
        pass

    try:
        texts = rt_ctrl.texts()
        if texts:
            tries.append("\n".join(texts))
    except Exception:
        pass

    try:
        tries.append(rt_ctrl.get_value())
    except Exception:
        pass

    try:
        tries.append(rt_ctrl.iface_value.CurrentValue)
    except Exception:
        pass

    best = ""
    for t in tries:
        if t and len(t.strip()) > len(best):
            best = t.strip()
    return best

def is_access_denied_error(exc):
    msg = str(exc).lower()
    return isinstance(exc, (COM_ERROR, PermissionError, OSError)) or "access is denied" in msg

def reacquire_context():
    main_win, rect = get_eladalah_window()
    tree_ctrl = find_tree_control(main_win)
    if tree_ctrl is None:
        raise RuntimeError("Could not find tv_Subjects tree control.")
    rt_ctrl = find_rt_txt(main_win)
    return main_win, tree_ctrl, rt_ctrl, rect

def robust_read_text(rt_ctrl, retries=2):
    empty_count = 0
    for _ in range(retries + 1):
        text = read_rt_txt(rt_ctrl)
        if text.strip():
            return text, empty_count
        empty_count += 1
        time.sleep(READ_WAIT_SECONDS)
    return "", empty_count

def find_tree_control(main_window):
    controls = main_window.descendants()
    for c in controls:
        auto_id = get_info(c, "automation_id")
        class_name = get_info(c, "class_name")
        control_type = get_info(c, "control_type")
        if (
            auto_id == "tv_Subjects"
            or "SysTreeView32" in class_name
            or control_type == "Tree"
        ):
            return c
    return None

def get_context():
    return reacquire_context()

def rect_center(rect):
    return ((rect.left + rect.right) // 2, (rect.top + rect.bottom) // 2)

def get_visible_tree_items(tree_ctrl):
    items = []
    tree_rect = tree_ctrl.rectangle()

    for item in tree_ctrl.descendants():
        if get_info(item, "control_type") != "TreeItem":
            continue
        try:
            if not item.is_visible():
                continue
        except Exception:
            pass

        try:
            r = item.rectangle()
            if r.bottom <= tree_rect.top or r.top >= tree_rect.bottom:
                continue
        except Exception:
            r = None

        name = get_name(item)
        items.append((item, name, r))

    items.sort(key=lambda x: (x[2].top if x[2] else 10**9, x[1]))
    return items

def visible_signature(visible_items):
    sig = []
    for item, name, r in visible_items:
        try:
            item_name = (get_name(item) or name or "").strip() or "(blank)"
            top = int((r.top if r else item.rectangle().top) / 20)
            sig.append((item_name, top))
        except Exception:
            pass
    return tuple(sig)

def click_tree_item(item, fallback_x=None):
    try:
        item.click_input()
        return True, "treeitem_click_input"
    except Exception:
        try:
            r = item.rectangle()
            x, y = rect_center(r)
            pyautogui.click(x, y)
            return True, "treeitem_rect_center_click"
        except Exception:
            if fallback_x is not None:
                try:
                    r = item.rectangle()
                    y = (r.top + r.bottom) // 2
                    pyautogui.click(fallback_x, y)
                    return True, "row_fallback_x_click"
                except Exception:
                    pass
    return False, "click_failed"

def try_expand_tree_item(item, tree_ctrl):
    before = visible_signature(get_visible_tree_items(tree_ctrl))

    def changed():
        after = visible_signature(get_visible_tree_items(tree_ctrl))
        return after != before

    try:
        exp = item.iface_expand_collapse
        state = exp.CurrentExpandCollapseState
        if state != 1:  # not Expanded
            exp.Expand()
            time.sleep(0.8)
            if changed():
                return True, "uia_expand"
    except Exception:
        pass

    try:
        item.click_input()
        time.sleep(0.2)
        pyautogui.press("right")
        time.sleep(0.8)
        if changed():
            return True, "right_arrow"
    except Exception:
        pass

    try:
        r = item.rectangle()
        pyautogui.doubleClick((r.left + r.right) // 2, (r.top + r.bottom) // 2)
        time.sleep(0.8)
        if changed():
            return True, "double_click"
    except Exception:
        pass

    try:
        r = item.rectangle()
        expander_x = max(r.left - 10, 0)
        expander_y = (r.top + r.bottom) // 2
        pyautogui.click(expander_x, expander_y)
        time.sleep(0.8)
        if changed():
            return True, "expander_click"
    except Exception:
        pass

    return False, ""

def try_move_tree_down(before_sig, main_win, tree_ctrl, rect):
    def after_changed():
        now_items = get_visible_tree_items(tree_ctrl)
        now_sig = visible_signature(now_items)
        return now_sig != before_sig, now_items, now_sig

    movement_methods = [
        ("PageDown", lambda: send_keys("{PGDN}")),
        ("Down*12", lambda: send_keys("{DOWN 12}")),
        ("Down*24", lambda: send_keys("{DOWN 24}")),
        ("MouseWheel", None),
    ]

    for method_name, action in movement_methods:
        focused, _ = safe_focus_window(main_win, rect)
        if not focused:
            continue
        time.sleep(0.2)
        cx, cy = None, None
        try:
            tree_rect = tree_ctrl.rectangle()
            cx, cy = rect_center(tree_rect)
            tree_ctrl.set_focus()
        except Exception:
            if cx is not None and cy is not None:
                pyautogui.click(cx, cy)
        time.sleep(0.15)

        if method_name == "MouseWheel":
            pyautogui.scroll(-700, x=cx, y=cy)
        else:
            action()
        time.sleep(MOVE_WAIT_SECONDS)

        changed, now_items, now_sig = after_changed()
        if changed:
            return True, method_name, now_items, now_sig

    _, now_items, now_sig = after_changed()
    return False, "none", now_items, now_sig

def load_done_hashes():
    done = set()

    if not LOG_FILE.exists():
        return done

    with LOG_FILE.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            if row.get("status") == "OK" and row.get("hash"):
                done.add(row["hash"])

    return done

def save_docx(text, docx_path, title):
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)

    for line in text.splitlines():
        line = line.strip()
        if line:
            doc.add_paragraph(line)

    doc.save(str(docx_path))

def main():
    pyautogui.FAILSAFE = True
    pyautogui.PAUSE = 0.2

    done_hashes = load_done_hashes()
    print(f"Already exported hashes: {len(done_hashes)}")

    main_win, tree_ctrl, rt_ctrl, _ = get_context()
    print("Tree control located.")

    file_exists = LOG_FILE.exists()

    with LOG_FILE.open("a", encoding="utf-8-sig", newline="") as f:
        fieldnames = [
            "status",
            "screen",
            "page",
            "visible_title",
            "row",
            "chars",
            "hash",
            "title",
            "docx_file",
            "txt_file",
            "error",
        ]
        writer = csv.DictWriter(f, fieldnames=fieldnames)

        if not file_exists or LOG_FILE.stat().st_size == 0:
            writer.writeheader()

        stagnant_pages = 0
        seen_page_signatures = set()
        repeat_page_count = 0

        page = 1
        while page <= MAX_PAGES:
            print(f"\n========== PAGE {page} ==========")

            main_win, tree_ctrl, rt_ctrl, rect = get_context()
            safe_focus_window(main_win, rect)
            time.sleep(0.2)
            visible_items = get_visible_tree_items(tree_ctrl)
            if MAX_ITEMS_PER_PAGE:
                visible_items = visible_items[:MAX_ITEMS_PER_PAGE]
            before_sig = visible_signature(visible_items)
            writer.writerow({
                "status": "PAGE_SEEN",
                "screen": page,
                "page": page,
                "visible_title": "",
                "row": 0,
                "chars": len(visible_items),
                "hash": "",
                "title": "",
                "docx_file": "",
                "txt_file": "",
                "error": f"action=PAGE_SEEN; signature_items={len(before_sig)}",
            })
            f.flush()

            if STOP_ON_REPEATED_PAGE_SIGNATURE and before_sig in seen_page_signatures:
                repeat_page_count += 1
                writer.writerow({
                    "status": "STOP_REPEATED_PAGE",
                    "screen": page,
                    "page": page,
                    "visible_title": "",
                    "row": 0,
                    "chars": len(visible_items),
                    "hash": "",
                    "title": "",
                    "docx_file": "",
                    "txt_file": "",
                    "error": (
                        "action=STOP_REPEATED_PAGE; "
                        f"repeat_count={repeat_page_count}/{STOP_AFTER_REPEAT_COUNT}"
                    ),
                })
                f.flush()
                print("Stopping: repeated visible page signature")
                if repeat_page_count >= STOP_AFTER_REPEAT_COUNT:
                    break
            else:
                seen_page_signatures.add(before_sig)
                repeat_page_count = 0

            before_names = [name or "(blank)" for _, name, _ in visible_items]
            print(f"Visible before move candidate ({len(before_names)}): {before_names}")

            new_this_page = 0
            expanded_any = False

            for idx, (item, item_name, _) in enumerate(visible_items, start=1):
                row_label = item_name or f"row-{idx}"

                try:
                    safe_focus_window(main_win, rect)
                    time.sleep(0.2)
                    current_items = get_visible_tree_items(tree_ctrl)
                    if idx > len(current_items):
                        continue
                    fresh_item, _, _ = current_items[idx - 1]
                    clicked, click_method = click_tree_item(fresh_item)
                    if not clicked:
                        raise RuntimeError("Failed to click tree item.")
                    time.sleep(CLICK_WAIT_SECONDS)

                    text, empty_reads = robust_read_text(rt_ctrl, retries=2 if FAST_MODE else 3)
                    if empty_reads >= 2:
                        main_win, tree_ctrl, rt_ctrl, rect = get_context()
                        text, _ = robust_read_text(rt_ctrl, retries=1)

                    if len(text.strip()) < MIN_CHARS:
                        print(f"SKIP_SHORT page={page} row={idx} title={row_label}")
                        writer.writerow({
                            "status": "SKIP_SHORT",
                            "screen": page,
                            "page": page,
                            "visible_title": row_label,
                            "row": idx,
                            "chars": len(text.strip()),
                            "hash": "",
                            "title": "",
                            "docx_file": "",
                            "txt_file": "",
                            "error": f"action=SKIP_SHORT; click={click_method}",
                        })
                        f.flush()
                    else:
                        h = text_hash(text)

                        if h in done_hashes:
                            print(f"DUPLICATE page={page} row={idx} title={row_label}")
                            writer.writerow({
                                "status": "DUPLICATE",
                                "screen": page,
                                "page": page,
                                "visible_title": row_label,
                                "row": idx,
                                "chars": len(text),
                                "hash": h,
                                "title": "",
                                "docx_file": "",
                                "txt_file": "",
                                "error": f"action=DUPLICATE; click={click_method}",
                            })
                            f.flush()
                        else:
                            lines = [x.strip() for x in text.splitlines() if x.strip()]
                            title = lines[0][:100] if lines else f"page-{page}-row-{idx}"
                            safe_title = clean_filename(title)

                            base_name = f"s{page:04d}-r{idx:02d}-{short_hash(text)}-{safe_title}"
                            docx_path = DOCX_DIR / f"{base_name}.docx"
                            txt_path = TXT_DIR / f"{base_name}.txt"

                            if SAVE_TXT:
                                txt_path.write_text(text, encoding="utf-8")
                            if SAVE_DOCX:
                                save_docx(text, docx_path, title)

                            done_hashes.add(h)
                            new_this_page += 1

                            print(f"EXPORTED page={page} row={idx} title={row_label} chars={len(text)}")
                            print(f"Saved: {docx_path.name}")

                            writer.writerow({
                                "status": "OK",
                                "screen": page,
                                "page": page,
                                "visible_title": row_label,
                                "row": idx,
                                "chars": len(text),
                                "hash": h,
                                "title": title,
                                "docx_file": str(docx_path) if SAVE_DOCX else "",
                                "txt_file": str(txt_path) if SAVE_TXT else "",
                                "error": f"action=EXPORTED; click={click_method}",
                            })
                            f.flush()

                    if AUTO_EXPAND:
                        expanded, method = try_expand_tree_item(fresh_item, tree_ctrl)
                        if expanded:
                            expanded_any = True
                            print(f"EXPANDED page={page} row={idx} title={row_label} method={method}")
                            writer.writerow({
                                "status": "EXPANDED",
                                "screen": page,
                                "page": page,
                                "visible_title": row_label,
                                "row": idx,
                                "chars": 0,
                                "hash": "",
                                "title": "",
                                "docx_file": "",
                                "txt_file": "",
                                "error": f"action=EXPANDED; expand_method={method}",
                            })
                            f.flush()
                            break

                    print(f"PROCESSED page={page} row={idx} title={row_label}")

                except Exception as e:
                    print(f"FAILED page={page} row={idx} title={row_label}")
                    print(traceback.format_exc())

                    writer.writerow({
                        "status": "FAILED",
                        "screen": page,
                        "page": page,
                        "visible_title": row_label,
                        "row": idx,
                        "chars": 0,
                        "hash": "",
                        "title": "",
                        "docx_file": "",
                        "txt_file": "",
                        "error": str(e),
                    })
                    f.flush()
                    if is_access_denied_error(e):
                        main_win, tree_ctrl, rt_ctrl, rect = get_context()

            if expanded_any:
                print(f"Restarting page {page} after expansion.")
                continue

            try:
                moved, nav_method, after_items, after_sig = try_move_tree_down(before_sig, main_win, tree_ctrl, rect)
            except Exception:
                main_win, tree_ctrl, rt_ctrl, rect = get_context()
                moved, nav_method, after_items, after_sig = try_move_tree_down(before_sig, main_win, tree_ctrl, rect)
            after_names = [name or "(blank)" for _, name, _ in after_items]
            print(f"Visible after move ({len(after_names)}): {after_names}")
            print(f"MOVE_DOWN page={page} moved={moved} method={nav_method}")
            print(f"New exports this page: {new_this_page}")

            if not moved:
                stagnant_pages += 1
            else:
                stagnant_pages = 0

            print(f"Stagnant pages: {stagnant_pages}/{STOP_AFTER_STAGNANT_PAGES}")
            if not moved:
                writer.writerow({
                    "status": "STAGNANT",
                    "screen": page,
                    "page": page,
                    "visible_title": "",
                    "row": 0,
                    "chars": 0,
                    "hash": "",
                    "title": "",
                    "docx_file": "",
                    "txt_file": "",
                    "error": "action=STAGNANT",
                })
            else:
                writer.writerow({
                    "status": "MOVE_DOWN",
                    "screen": page,
                    "page": page,
                    "visible_title": "",
                    "row": 0,
                    "chars": 0,
                    "hash": "",
                    "title": "",
                    "docx_file": "",
                    "txt_file": "",
                    "error": f"action=MOVE_DOWN; moved={moved}; method={nav_method}",
                })
            f.flush()

            if stagnant_pages >= STOP_AFTER_STAGNANT_PAGES:
                print("Stopping: repeated navigation attempts produced no new items/hashes.")
                break
            page += 1

    print("\nFull export finished.")
    if SAVE_DOCX:
        print(f"DOCX: {DOCX_DIR}")
    if SAVE_TXT:
        print(f"TXT:  {TXT_DIR}")
    print(f"LOG:  {LOG_FILE}")

if __name__ == "__main__":
    main()
